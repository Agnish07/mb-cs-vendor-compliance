from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Vendor_Compliance_Interview_Architecture_Guide.docx"


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(header_cells[idx], header, bold=True)
        shade_cell(header_cells[idx], "D9EAD3")

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)

    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)

    doc.add_paragraph()


def add_code_block(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = "Code"
    run = paragraph.add_run(text)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(8)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    shade_cell(cell, "EAF3F8")
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10)
    p.add_run(f"\n{body}").font.size = Pt(9)
    doc.add_paragraph()


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)

    for style_name, size, color in [
        ("Title", 24, "1F4E5F"),
        ("Heading 1", 18, "1F4E5F"),
        ("Heading 2", 14, "2D7A78"),
        ("Heading 3", 12, "404040"),
    ]:
        style = styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    if "Code" not in [style.name for style in styles]:
        styles.add_style("Code", 1)
    styles["Code"].font.name = "Courier New"
    styles["Code"].font.size = Pt(8)


def build_document() -> None:
    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Autonomous Vendor Compliance System")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(31, 78, 95)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = subtitle.add_run("Architecture, Implementation, AI Agents, Demo Script, and Interview Preparation Guide")
    sub.font.size = Pt(12)
    sub.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()
    add_callout(
        doc,
        "Best one-line positioning",
        "This project is a production-style autonomous compliance platform that ingests vendor duty files through email, normalizes heterogeneous formats, validates customs rules, uses Gemini-powered agents for triage/remediation/insights, and gives compliance teams a live governance dashboard.",
    )

    doc.add_heading("1. Executive Interview Pitch", level=1)
    doc.add_heading("30-second answer", level=2)
    doc.add_paragraph(
        "I built an autonomous vendor compliance system for part-level customs duty governance. Vendors interact only through email, so the platform reads attachments from Gmail, extracts CSV/XLSX/PDF files, normalizes them into a canonical customs schema, validates HSN/RITC and duty components, classifies errors as isolated, pattern, or systemic, sends correction emails, tracks escalations, computes vendor risk, and displays everything in a React dashboard. I combined deterministic rule-based validation for auditability with Gemini-powered agents for fuzzy tasks like column mapping, PDF extraction fallback, triage, remediation emails, audit narratives, and trend insights."
    )

    doc.add_heading("2-minute answer", level=2)
    doc.add_paragraph(
        "The case study has two core problems: messy vendor-submitted customs data and weak vendor governance. My architecture solves this by creating a closed-loop compliance automation layer. The backend is FastAPI with SQLAlchemy and PostgreSQL. It has services that act like agents: ingestion, normalization, validation, error classification, remediation, escalation, risk, and trend mining. The frontend is a Vite React dashboard using Axios and Recharts. The AI layer uses Gemini, but the system does not blindly trust AI for correctness. The compliance rules are deterministic; AI is used where ambiguity is unavoidable, such as mapping unknown columns, extracting PDF rows, generating operational triage recommendations, and writing vendor-facing emails."
    )

    doc.add_heading("How to say the design principle", level=2)
    add_bullets(
        doc,
        [
            "Rules decide whether a submission is compliant.",
            "AI explains, maps, extracts, triages, and communicates.",
            "Every incoming file runs through the same pipeline, whether uploaded manually or received by email.",
            "Every error becomes structured data, not just a log line.",
            "The dashboard is not cosmetic; it is the governance surface for risk, trends, escalations, correct entries, and audit history.",
        ]
    )

    doc.add_heading("2. Problem-to-Solution Mapping", level=1)
    add_table(
        doc,
        ["Case study problem", "Implemented system response", "Why this matters"],
        [
            [
                "Vendors submit heterogeneous, non-standardized files.",
                "Normalization service maps CSV/XLSX/PDF inputs into one canonical schema.",
                "Downstream validation and analytics become consistent regardless of vendor format.",
            ],
            [
                "Manual deconvolution into part identity, HSN/RITC, and duty vectors.",
                "Pipeline extracts part_number, description, hsn_code, bcd, cvd, sws, and igst automatically.",
                "Reduces manual effort and creates repeatable controls.",
            ],
            [
                "Missing attributes and invalid master references.",
                "Validation service checks mandatory fields, HSN master list, missing duty components, and duty ranges.",
                "Prevents bad data from reaching duty computation or compliance reporting.",
            ],
            [
                "Weak vendor governance and no objective measurement.",
                "Vendor risk service computes row-failure rate, severity, recurrence, systemic failures, and score.",
                "Creates measurable vendor performance signals.",
            ],
            [
                "No escalation logic for delayed corrections.",
                "Escalation service tracks correction_due_at and increments escalation_level for overdue vendors.",
                "Turns correction follow-up into an auditable process.",
            ],
            [
                "Email-only vendor interaction.",
                "IMAP ingestion reads files; SMTP remediation sends correction/escalation emails.",
                "Matches the constraint of no approved vendor portal/interface.",
            ],
            [
                "Need to discriminate isolated mistakes from systemic failure modes.",
                "Classification service labels ISOLATED, PATTERN, and SYSTEMIC errors.",
                "Helps prioritize vendor governance action.",
            ],
        ],
        widths=[2.2, 2.8, 2.6],
    )

    doc.add_heading("3. Tech Stack", level=1)
    add_table(
        doc,
        ["Layer", "Technology", "Reason"],
        [
            ["Backend API", "FastAPI", "Fast async-friendly API framework with Pydantic response models and automatic OpenAPI docs."],
            ["ORM", "SQLAlchemy", "Clean database modeling and query abstraction."],
            ["Database", "PostgreSQL", "Production-grade relational store for submissions, errors, rows, audit state, and risk history."],
            ["Configuration", "pydantic-settings + .env", "Keeps environment-specific settings outside code."],
            ["Email ingestion", "IMAP", "Reads unread vendor emails and attachments from Gmail."],
            ["Email sending", "SMTP", "Sends correction and escalation emails."],
            ["AI", "Google Gemini API", "Handles ambiguous extraction, triage, language generation, and insight generation."],
            ["Frontend", "React + Vite", "Fast dashboard development and local dev experience."],
            ["HTTP client", "Axios", "Simple API integration from frontend to FastAPI."],
            ["Charts", "Recharts", "React-native charting for trajectory, distribution, and vendor analytics."],
        ],
        widths=[1.5, 2.0, 4.2],
    )

    doc.add_heading("4. High-Level Architecture", level=1)
    add_code_block(
        doc,
        """Vendor Gmail Inbox
   |
   | IMAP polling / manual upload
   v
FastAPI Backend
   |
   +--> Ingestion Agent: fetch unread emails, extract attachments, identify vendor
   +--> Normalization Agent: CSV/XLSX/PDF to canonical rows
   +--> Validation Agent: mandatory fields, HSN master, duty range, missing duties
   +--> Classification Agent: ISOLATED / PATTERN / SYSTEMIC
   +--> Gemini Triage Agent: operational category + recommendation
   +--> Remediation Agent: correction email through SMTP
   +--> Escalation Agent: correction SLA tracking
   +--> Risk + Trend Agent: score, trajectory, forecast, AI insight
   |
   v
PostgreSQL Database
   |
   v
React Dashboard
   |
   +--> Submissions, Errors, Correct Entries
   +--> Vendor Risk Panel
   +--> Escalation Governance
   +--> Recharts Analytics
"""
    )

    doc.add_heading("Architecture in interview language", level=2)
    doc.add_paragraph(
        "The architecture is event-driven at the business level and service-oriented inside the backend. A vendor email or manual upload triggers the same processing pipeline. Each pipeline stage converts unstructured or semi-structured vendor input into increasingly structured compliance data. The database becomes the source of truth, and the dashboard reads aggregated views from API endpoints."
    )

    doc.add_heading("5. Low-Level Backend Architecture", level=1)
    add_table(
        doc,
        ["Module", "Responsibility", "Interview explanation"],
        [
            ["app/main.py", "Initializes FastAPI, CORS, DB startup, routers, and optional polling task.", "This is the application entrypoint and lifecycle coordinator."],
            ["app/config.py", "Loads .env configuration.", "Keeps secrets, ports, polling, DB URL, Gmail, SMTP, and Gemini settings externalized."],
            ["app/database.py", "SQLAlchemy engine/session setup and initialization.", "Provides database sessions to routers and services."],
            ["models/compliance.py", "Submission, SubmissionRow, ComplianceError ORM models.", "Represents audit history, normalized row data, and validation exceptions."],
            ["routers/compliance.py", "REST endpoints for upload, submissions, errors, risks, trends, escalations, correct entries.", "Thin routing layer that delegates business logic to services."],
            ["services/ingestion.py", "Runs email attachment ingestion.", "Connects email intake with the same processing pipeline."],
            ["services/normalization.py", "Reads CSV/XLSX/PDF and maps to canonical schema.", "Solves heterogeneous vendor input."],
            ["services/validation.py", "Rule-based compliance validation.", "Deterministic control layer for auditability."],
            ["services/classification.py", "Error type classification.", "Distinguishes isolated, repeated, and systemic failures."],
            ["services/ai_agent.py", "Gemini-powered agent methods.", "Centralized AI layer with constrained prompts and fallbacks."],
            ["services/remediation.py", "Correction email workflow.", "Creates closed-loop vendor remediation."],
            ["services/escalation.py", "SLA breach detection and escalation.", "Automates governance follow-up."],
            ["services/risk.py", "Vendor risk score computation.", "Quantifies vendor quality and recurring failure risk."],
            ["services/trends.py", "Trend mining and AI insight generation.", "Turns historical failures into operational signals."],
        ],
        widths=[2.0, 2.7, 3.0],
    )

    doc.add_heading("6. Frontend Architecture", level=1)
    doc.add_paragraph(
        "The frontend is a single-page React dashboard. It calls the FastAPI backend through Axios, keeps the dashboard state in React hooks, and renders metrics, tables, risk cards, escalation status, and charts with Recharts."
    )
    add_table(
        doc,
        ["Frontend part", "Role"],
        [
            ["src/services/api.js", "Axios client configured with VITE_API_BASE_URL and dashboard API calls."],
            ["src/App.jsx", "Main dashboard state, API loading, upload form, metrics, charts, tables, and trend display."],
            ["src/App.css", "Dashboard layout, cards, risk colors, responsive behavior, and table styling."],
            ["Recharts", "LineChart for row-failure trajectory, PieChart for error distribution, BarChart for vendor and error-code analytics."],
        ],
        widths=[2.2, 5.5],
    )

    doc.add_heading("7. Database Design", level=1)
    doc.add_heading("Core tables", level=2)
    add_table(
        doc,
        ["Table", "Important columns", "Purpose"],
        [
            [
                "submissions",
                "id, vendor_email, file_name, upload_time, total_rows, error_count, status, correction_due_at, remediation_sent_at, escalation_level, resolved_at, ai_audit_summary",
                "One record per received file. Tracks lifecycle, SLA, and audit narrative.",
            ],
            [
                "submission_rows",
                "submission_id, row_number, part_number, description, hsn_code, bcd, cvd, sws, igst, is_valid, error_summary",
                "Stores normalized row-level data and whether each row passed validation.",
            ],
            [
                "errors",
                "submission_id, row_number, error_code, severity, error_type, message, triage_category, ai_recommendation",
                "Stores every validation exception in structured form.",
            ],
        ],
        widths=[1.5, 3.5, 2.7],
    )

    doc.add_heading("Why this model is interview-friendly", level=2)
    add_bullets(
        doc,
        [
            "It separates file-level history from row-level normalized data and error-level exception details.",
            "It supports auditability because every error is stored with row number, code, severity, type, and recommendation.",
            "It supports analytics because vendor risk and trends can be computed from historical submissions and errors.",
            "It supports correct-entry visibility because valid rows are stored, not discarded.",
        ]
    )

    doc.add_heading("8. End-to-End Processing Flow", level=1)
    doc.add_heading("Manual upload flow", level=2)
    add_numbered(
        doc,
        [
            "User uploads a CSV/XLSX/PDF from the dashboard using POST /upload.",
            "FastAPI receives the file and vendor_email.",
            "Pipeline saves the file locally.",
            "Normalization converts the file into canonical rows.",
            "Validation checks mandatory fields, HSN/RITC master reference, duty range, and missing duty components.",
            "Classification labels errors as ISOLATED, PATTERN, or SYSTEMIC.",
            "Gemini triages errors and produces recommendations.",
            "Submission, rows, and errors are stored in the database.",
            "If errors exist, remediation email is generated and sent if email automation is enabled.",
            "Dashboard refreshes and shows updated risk, trends, errors, submissions, and correct entries.",
        ]
    )

    doc.add_heading("Email ingestion flow", level=2)
    add_numbered(
        doc,
        [
            "Email polling starts on FastAPI startup if EMAIL_POLLING_ENABLED is true.",
            "Email service connects to Gmail through IMAP.",
            "Unread messages are scanned for supported attachments: CSV, XLS, XLSX, and PDF.",
            "Sender email becomes the vendor identity.",
            "Attachments are saved locally.",
            "Each attachment is sent through the same processing pipeline as manual upload.",
            "If a submission has errors, the system sends a structured correction email through SMTP.",
            "Vendor replies with corrected file; the email poller processes that attachment as a new submission from the same vendor.",
        ]
    )

    doc.add_heading("9. Canonical Schema and Normalization", level=1)
    doc.add_paragraph("Every vendor file is normalized into the same schema:")
    add_code_block(
        doc,
        """part_number
description
hsn_code
bcd
cvd
sws
igst"""
    )
    doc.add_paragraph(
        "The normalization layer handles inconsistent column names. Deterministic mapping covers known variations, and Gemini can infer mappings for unknown headers. For PDFs, deterministic extraction is attempted first. If the PDF text is difficult to parse, Gemini receives the extracted text and returns structured JSON rows."
    )

    doc.add_heading("10. Validation Rules", level=1)
    add_table(
        doc,
        ["Rule", "Example", "Error code", "Why it matters"],
        [
            ["Missing mandatory field", "part_number or hsn_code is blank", "MISSING_MANDATORY_FIELD", "Cannot identify part or classify duty correctly."],
            ["Invalid HSN/RITC", "HSN not found in master list", "INVALID_HSN_CODE", "Invalid master reference can cause regulatory and duty errors."],
            ["Duty out of range", "IGST = 140 or BCD = -5", "DUTY_OUT_OF_RANGE", "Duty rates outside 0-100 indicate invalid computation input."],
            ["Missing duty component", "BCD/CVD/SWS/IGST missing", "MISSING_DUTY_COMPONENT", "Incomplete duty vector prevents correct landed-cost/duty calculation."],
        ],
        widths=[2.0, 2.0, 2.0, 2.2],
    )

    doc.add_heading("11. Error Classification", level=1)
    doc.add_paragraph("After validation, the system classifies error patterns.")
    add_table(
        doc,
        ["Type", "Meaning", "Rule"],
        [
            ["ISOLATED", "A one-off row-level issue.", "Single occurrence or not repeated enough."],
            ["PATTERN", "Repeated issue within the same file.", "Same error appears at least 3 times in a submission."],
            ["SYSTEMIC", "Repeated across vendor submissions.", "Same vendor repeats the same failure across multiple submissions."],
        ],
        widths=[1.5, 3.2, 3.0],
    )
    add_callout(
        doc,
        "Strong interview point",
        "This is where the project directly addresses the case-study requirement to discriminate isolated mistakes from systemic vendor failure modes.",
    )

    doc.add_heading("12. Agentic AI Implementation", level=1)
    doc.add_paragraph(
        "Agentic AI is implemented as specialized Gemini-backed methods inside GeminiComplianceAgent. The system is agentic because each AI capability has a role, structured context, task-specific prompt, constrained output, and deterministic fallback."
    )
    add_table(
        doc,
        ["AI agent", "Input context", "Output", "Why AI is useful"],
        [
            ["Column Mapping Agent", "Source columns, canonical schema, existing deterministic mapping.", "JSON mapping from canonical fields to vendor columns.", "Column names vary widely; AI handles semantic similarity."],
            ["PDF Extraction Agent", "Extracted PDF text and canonical fields.", "Structured rows JSON.", "PDF tables can be irregular and hard to parse deterministically."],
            ["Exception Triage Agent", "Structured validation errors.", "Triage category and recommendation.", "Converts technical errors into operational categories."],
            ["Remediation Email Agent", "Vendor email, file name, error rows.", "Vendor-facing correction email.", "Generates clear, professional remediation instructions."],
            ["Audit Narrative Agent", "Submission context and result.", "Audit-ready summary.", "Explains what happened in compliance language."],
            ["Escalation Email Agent", "Overdue submission context.", "Escalation email.", "Creates firm but professional follow-up communication."],
            ["Vendor Risk Insight Agent", "Vendor risk metrics.", "Short risk explanation.", "Turns score into business language."],
            ["Trend Mining Insight Agent", "Failure rate, forecast, risk band, trajectory, dominant errors, severity/triage mix.", "Dashboard trend insight.", "Explains charts and suggests governance action."],
        ],
        widths=[1.8, 2.3, 1.8, 2.0],
    )

    doc.add_heading("Important AI safety/design decision", level=2)
    add_bullets(
        doc,
        [
            "The system does not let AI decide final compliance correctness.",
            "Validation remains rule-based and deterministic.",
            "AI is used for fuzzy interpretation, natural-language generation, and operational recommendations.",
            "Every AI function has a fallback so the pipeline does not break if Gemini is unavailable.",
            "JSON-returning prompts are parsed and validated before being trusted.",
        ]
    )

    doc.add_heading("13. Remediation and Closed-Loop Email Automation", level=1)
    doc.add_paragraph(
        "When validation errors exist, the submission becomes PENDING_CORRECTION. The remediation agent prepares a structured correction email. If email automation is enabled, SMTP sends the email to the vendor. The vendor can reply with a corrected attachment, which is picked up by the IMAP ingestion cycle and reprocessed through the same validation pipeline."
    )
    add_code_block(
        doc,
        """Bad vendor file
   -> errors detected
   -> correction email sent
   -> vendor replies with corrected attachment
   -> attachment reprocessed
   -> accepted if error_count == 0, otherwise another correction loop"""
    )
    add_callout(
        doc,
        "Honest production note",
        "In the current prototype, corrected files are tracked as new submissions from the same vendor. A production extension would link corrections to the original submission using Gmail thread ID, message ID, or a remediation token in the subject/body.",
    )

    doc.add_heading("14. Escalation Governance", level=1)
    doc.add_paragraph(
        "The escalation agent solves the governance weakness mentioned in the case study. It tracks correction_due_at. If a submission remains PENDING_CORRECTION after its SLA, the system marks it ESCALATED, increments escalation_level, optionally sends an escalation email, and displays it in the dashboard."
    )
    add_code_block(
        doc,
        """if submission.status == "PENDING_CORRECTION" and submission.correction_due_at < now:
    submission.status = "ESCALATED"
    submission.escalation_level += 1
    Gemini drafts escalation email
    SMTP sends follow-up
    dashboard shows open escalation"""
    )

    doc.add_heading("15. Vendor Risk Scoring", level=1)
    doc.add_paragraph(
        "The risk model is designed to avoid misleading percentages above 100%. It uses unique failed rows instead of raw error count for error rate. One row can have multiple validation errors, but it is still one failed row for row-failure-rate purposes."
    )
    add_code_block(
        doc,
        """row_failure_rate = unique_failed_rows / total_rows
risk_score = 100
  - quality_penalty
  - severity_penalty
  - systemic_repeat_penalty
  - recurrence_penalty"""
    )
    add_table(
        doc,
        ["Signal", "Meaning"],
        [
            ["Quality penalty", "How much of the vendor's submitted row population failed validation."],
            ["Severity penalty", "How serious the errors are, normalized against maximum possible severity."],
            ["Systemic repeat penalty", "Penalty for repeated systemic errors across submissions."],
            ["Recurrence penalty", "Penalty when many submissions from the vendor contain errors."],
        ],
        widths=[2.2, 5.5],
    )

    doc.add_heading("16. Trend Mining", level=1)
    doc.add_paragraph(
        "Trend mining computes vendor trajectories and forecast-like signals from historical submissions. It calculates row-failure rate per submission, baseline rate, recent rate, slope, next-submission forecast, risk band, dominant error code, triage mix, severity mix, and a Gemini-generated dashboard insight."
    )
    add_table(
        doc,
        ["Metric", "Meaning"],
        [
            ["Trajectory", "Direction of movement: IMPROVING, STABLE, or DETERIORATING."],
            ["Risk band", "Current risk severity: LOW, MEDIUM, HIGH."],
            ["Recent row-failure rate", "Average failed-row percentage in recent submissions."],
            ["Forecast next rate", "Simple slope-based projection of next row-failure rate, capped at 100%."],
            ["Dominant error code", "Most frequent error driving vendor risk."],
        ],
        widths=[2.2, 5.5],
    )
    add_callout(
        doc,
        "How to explain STABLE + HIGH RISK",
        "STABLE describes direction, not safety. A vendor can be stable because its error rate is consistently bad. That is why the dashboard separates trajectory from risk band.",
    )

    doc.add_heading("17. API Reference", level=1)
    add_table(
        doc,
        ["Endpoint", "Method", "Purpose"],
        [
            ["/health", "GET", "Health check."],
            ["/submissions", "GET", "Submission history."],
            ["/errors", "GET", "Validation error logs."],
            ["/correct-entries", "GET", "All rows that passed validation."],
            ["/vendor-risk", "GET", "Vendor risk scores and AI risk insights."],
            ["/trend-insights", "GET", "Trend analytics and Gemini-generated trend narratives."],
            ["/escalations", "GET", "Open correction/escalation items."],
            ["/run-escalations", "POST", "Manually trigger escalation cycle."],
            ["/upload", "POST", "Manual file intake through dashboard."],
            ["/ingest-emails", "POST", "Manual trigger for unread email attachment ingestion."],
        ],
        widths=[2.0, 1.0, 4.7],
    )

    doc.add_heading("18. Graph Implementation", level=1)
    add_table(
        doc,
        ["Graph", "Component", "Data source", "Meaning"],
        [
            ["Error-rate trajectory", "Recharts LineChart", "/trend-insights timeline", "Submission-by-submission row-failure rate."],
            ["Error distribution", "Recharts PieChart", "/errors", "Distribution of validation errors by error_code."],
            ["Errors per vendor", "Recharts BarChart", "/submissions", "Total errors accumulated per vendor."],
            ["Recurring error codes", "Recharts BarChart", "/trend-insights top_error_codes", "Most frequent recurring validation issues."],
        ],
        widths=[2.0, 1.7, 2.1, 2.0],
    )

    doc.add_heading("19. Demo Script", level=1)
    doc.add_heading("Before demo", level=2)
    add_numbered(
        doc,
        [
            "Start PostgreSQL if using local Postgres.",
            "Start backend: uvicorn app.main:app --reload --port 8010 from BACKEND.",
            "Start frontend: npm run dev from frontend.",
            "Open dashboard at http://localhost:5173.",
            "Confirm .env has database, Gmail, SMTP, Gemini, and polling settings.",
        ]
    )
    doc.add_heading("Demo flow", level=2)
    add_numbered(
        doc,
        [
            "Show dashboard overview: submissions, errors, vendors, rows inspected, correct entries.",
            "Upload a mixed-validity CSV/PDF file.",
            "Explain normalization into canonical schema.",
            "Show submission history and error logs.",
            "Show correct entries section to prove valid rows are retained.",
            "Show vendor risk panel and explain score components.",
            "Show trend mining and explain Gemini-generated insight.",
            "Show escalation governance and explain correction SLA tracking.",
            "Optionally send an email with an attachment and show that the IMAP agent picks it up.",
        ]
    )
    add_heading_text = "Strong demo narration"
    add_callout(
        doc,
        add_heading_text,
        "The key thing to emphasize is that the same compliance pipeline runs for both dashboard upload and email ingestion. That makes the system consistent, testable, and easy to extend.",
    )

    doc.add_heading("20. Interview Questions and Strong Answers", level=1)
    qa = [
        (
            "Why did you combine rule-based validation with AI?",
            "Compliance correctness needs deterministic, auditable rules. AI is used for ambiguity and communication: messy columns, PDFs, triage recommendations, emails, audit narratives, and insights. This gives the system both reliability and intelligence.",
        ),
        (
            "How is this agentic AI?",
            "Each agent has a specialized role, structured context, constrained prompt, output contract, and fallback. The agents operate around the pipeline: mapping, extraction, triage, remediation, audit, escalation, risk, and trend insight.",
        ),
        (
            "What happens if Gemini fails?",
            "The system uses deterministic fallbacks. Validation and storage still work. AI-generated recommendations or narratives fall back to template-based content.",
        ),
        (
            "Why is row-failure rate better than error_count / rows?",
            "A single row can produce multiple validation errors. Using raw errors can exceed 100%. Row-failure rate measures how much of the file is affected and stays interpretable.",
        ),
        (
            "How are corrections verified?",
            "The corrected attachment is reprocessed through the same normalization and validation pipeline. If no errors remain, it is accepted; otherwise, the remediation loop continues.",
        ),
        (
            "How would you productionize this?",
            "I would add Alembic migrations, authentication/RBAC, Celery or a worker queue for email polling and AI jobs, retry/dead-letter handling, structured audit logs, Docker deployment, monitoring, test coverage, and Gmail OAuth.",
        ),
        (
            "How do you prevent AI hallucination?",
            "Prompts explicitly say not to invent facts. The AI receives structured context and constrained outputs. For JSON outputs, responses are parsed and validated. AI does not decide final compliance pass/fail.",
        ),
        (
            "What is the most important design tradeoff?",
            "The tradeoff is using AI for fuzzy enterprise workflows without compromising deterministic compliance control. I solved it by placing rule validation at the center and AI around the edges.",
        ),
    ]
    for question, answer in qa:
        doc.add_heading(question, level=2)
        doc.add_paragraph(answer)

    doc.add_heading("21. Production Roadmap", level=1)
    add_table(
        doc,
        ["Area", "Upgrade", "Why"],
        [
            ["Database", "Alembic migrations", "Safe schema evolution across environments."],
            ["Security", "Auth, RBAC, secret manager", "Protect vendor/compliance data and credentials."],
            ["Email", "OAuth and Gmail API", "Better enterprise security than app passwords."],
            ["Background jobs", "Celery/RQ/APScheduler", "Reliable polling, retries, and long-running task management."],
            ["Audit", "Dedicated audit_events table", "Immutable trace of all decisions and communications."],
            ["Observability", "Structured logs, metrics, alerts", "Operational visibility and incident response."],
            ["Testing", "Unit/integration/e2e tests", "Confidence in validation, ingestion, and AI fallback behavior."],
            ["Deployment", "Docker Compose / Kubernetes", "Repeatable environment setup and scaling."],
            ["Data governance", "Master-data sync for HSN/RITC", "Replace mock master list with governed reference data."],
        ],
        widths=[1.5, 2.4, 3.8],
    )

    doc.add_heading("22. Resume/LinkedIn Bullet Points", level=1)
    add_bullets(
        doc,
        [
            "Built an autonomous vendor compliance platform using FastAPI, SQLAlchemy, PostgreSQL, React, Recharts, Gmail IMAP/SMTP, and Gemini API.",
            "Implemented multi-agent AI workflows for column mapping, PDF extraction fallback, exception triage, remediation email generation, audit narratives, risk insights, and trend mining.",
            "Designed a deterministic validation engine for HSN/RITC references, mandatory attributes, duty components, and duty range checks.",
            "Created a closed-loop email remediation workflow that ingests vendor files, sends correction requests, reprocesses replies, and tracks escalation status.",
            "Built a React governance dashboard with vendor risk scoring, trend analytics, correct-entry history, escalation tracking, and validation error logs.",
        ]
    )

    doc.add_heading("23. Final Interview Close", level=1)
    add_callout(
        doc,
        "Closing statement",
        "What I like about this project is that it is not just an AI wrapper. The core compliance decisions are deterministic and auditable, while AI improves the parts that are naturally messy: vendor formats, PDFs, triage language, remediation communication, and insight generation. That balance is what makes it suitable for enterprise compliance workflows.",
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
